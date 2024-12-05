from pathlib import Path
import logging
from typing import Dict, Any, Optional
from docx import Document

class DocumentGenerator:
    """Document generator for various formats"""

    def __init__(self, config: Dict[str, Any]):
        self.config = config

    def generate(self, content: Dict[str, Any], output_format: str, output_path: str) -> str:
        """Generate document in specified format"""
        try:
            if output_format not in ['pdf', 'docx', 'markdown']:
                raise ValueError(f"Unsupported format: {output_format}")

            if output_format == 'docx':
                return self._generate_docx(content, output_path)
            elif output_format == 'pdf':
                return self._generate_pdf(content, output_path)
            else:  # markdown
                return self._generate_markdown(content, output_path)
        except Exception as e:
            logging.error(f"Document generation failed: {str(e)}")
            raise

    def _generate_docx(self, content: Dict[str, Any], output_path: str) -> str:
        """Generate Word document"""
        doc = Document()

        # Apply formatting from config
        formatting = self.config.get('formatting', {})
        if formatting:
            doc.styles['Normal'].font.name = formatting.get('font', 'Times New Roman')
            doc.styles['Normal'].font.size = formatting.get('size', 12)

        # Add title
        if 'title' in content:
            doc.add_heading(content['title'], 0)

        # Add content
        if 'content' in content:
            doc.add_paragraph(content['content'])

        doc.save(output_path)
        return output_path

    def _generate_pdf(self, content: Dict[str, Any], output_path: str) -> str:
        """Generate PDF document"""
        # For now, we'll just create a markdown file
        return self._generate_markdown(content, output_path.replace('.pdf', '.md'))

    def _generate_markdown(self, content: Dict[str, Any], output_path: str) -> str:
        """Generate Markdown document"""
        markdown_content = f"# {content.get('title', '')}\n\n{content.get('content', '')}"
        
        with open(output_path, 'w') as f:
            f.write(markdown_content)
        return output_path