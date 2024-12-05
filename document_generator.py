from typing import Dict, Any, Optional
import subprocess
import os
from pathlib import Path
import jinja2
import markdown
from docx import Document
import logging
from dataclasses import dataclass

@dataclass
class DocumentFormat:
    """Document format specification"""
    extension: str
    mime_type: str
    template_type: str
    
    @classmethod
    def from_enum(cls, format_enum: OutputFormat) -> 'DocumentFormat':
        """Create DocumentFormat from OutputFormat enum"""
        format_map = {
            OutputFormat.MARKDOWN: cls('md', 'text/markdown', 'markdown'),
            OutputFormat.LATEX: cls('pdf', 'application/pdf', 'latex'),
            OutputFormat.PLAIN: cls('txt', 'text/plain', 'plain')
        }
        return format_map[format_enum]

class DocumentGenerator:
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.logger = logging.getLogger(__name__)
        
        # Initialize template engine
        self.template_env = jinja2.Environment(
            loader=jinja2.FileSystemLoader('templates'),
            trim_blocks=True,
            lstrip_blocks=True
        )
        
        # Define supported formats
        self.formats = {
            'pdf': DocumentFormat('pdf', 'application/pdf', 'latex'),
            'docx': DocumentFormat('docx', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'docx'),
            'markdown': DocumentFormat('md', 'text/markdown', 'markdown')
        }
        
    def generate_document(self, 
                         content: Dict[str, Any], 
                         output_format: str, 
                         output_path: str) -> str:
        """Generate document in specified format"""
        try:
            if output_format not in self.formats:
                raise ValueError(f"Unsupported format: {output_format}")
                
            # Prepare content with template
            formatted_content = self._apply_template(content)
            
            # Generate document based on format
            if output_format == 'pdf':
                return self._generate_pdf(formatted_content, output_path)
            elif output_format == 'docx':
                return self._generate_docx(formatted_content, output_path)
            else:  # markdown
                return self._generate_markdown(formatted_content, output_path)
                
        except Exception as e:
            self.logger.error(f"Document generation failed: {str(e)}")
            raise
            
    def _apply_template(self, content: Dict[str, Any]) -> str:
        """Apply academic template to content"""
        template = self.template_env.get_template(
            f"{self.config['template_variables']['TEMPLATE']['format']}.tex"
        )
        
        return template.render(
            title=content.get('title', ''),
            author=content.get('author', ''),
            sections=content.get('sections', []),
            bibliography=content.get('bibliography', []),
            formatting=self.config['template_variables']['TEMPLATE']['formatting']
        )
        
    def _generate_pdf(self, content: str, output_path: str) -> str:
        """Generate PDF using XeLaTeX"""
        output_dir = Path(output_path).parent
        temp_dir = output_dir / "temp"
        temp_dir.mkdir(exist_ok=True)
        
        tex_path = temp_dir / f"{Path(output_path).stem}.tex"
        
        try:
            # Check if xelatex is installed
            subprocess.run(['xelatex', '--version'], 
                          check=True, 
                          capture_output=True)
        except subprocess.CalledProcessError:
            raise RuntimeError("XeLaTeX is not installed or not in PATH")
        
        try:
            # Write TEX file
            with open(tex_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            # Run XeLaTeX twice for references
            for _ in range(2):
                subprocess.run([
                    'xelatex',
                    '-interaction=nonstopmode',
                    '-output-directory=' + str(output_dir),
                    str(tex_path)
                ], check=True, capture_output=True)
            
            return output_path
        
        except Exception as e:
            self.logger.error(f"PDF generation failed: {str(e)}")
            raise
        finally:
            # Cleanup temporary files
            if temp_dir.exists():
                import shutil
                shutil.rmtree(temp_dir)
            
    def _generate_docx(self, content: str, output_path: str) -> str:
        """Generate Word document"""
        doc = Document()
        
        # Apply formatting
        formatting = self.config['template_variables']['TEMPLATE']['formatting']
        doc.styles['Normal'].font.name = formatting['font']
        doc.styles['Normal'].font.size = formatting['size']
        
        # Add content
        for section in content.get('sections', []):
            doc.add_heading(section['title'], level=1)
            doc.add_paragraph(section['content'])
            
        # Save document
        doc.save(output_path)
        return output_path
        
    def _generate_markdown(self, content: str, output_path: str) -> str:
        """Generate Markdown document"""
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        return output_path 