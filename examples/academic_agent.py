"""
AgentFlow Academic Research Workflow Example

This example demonstrates how to create a comprehensive academic research workflow
using ELL2A's integration capabilities. The workflow includes:
1. Literature Review
2. Research Methodology
3. Data Analysis
4. Paper Writing

Requirements:
- OpenAI API key set in environment variable OPENAI_API_KEY
"""

import asyncio
import os
import json
from typing import List, Dict, Any, Optional
from pydantic import BaseModel, Field
import openai
from docx import Document
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

from agentflow.ell2a.integration import ELL2AIntegration
from agentflow.ell2a.types.message import Message, MessageRole

# Get singleton instance
ell2a = ELL2AIntegration()

def get_env_var(name: str) -> str:
    """Get an environment variable or raise an error if it's not set."""
    value = os.getenv(name)
    if not value:
        raise ValueError(f"Environment variable {name} is not set")
    return value

try:
    # Get Deepseek API key from environment variable or use the one from config
    api_key = os.getenv("DEEPSEEK_API_KEY", "sk-3671f0e962bf4e4980750a47c549ccd4")
    
    # Create OpenAI client with Deepseek configuration
    client = openai.OpenAI(
        api_key=api_key,
        base_url="https://api.deepseek.com/v1"  # Deepseek API endpoint
    )

except ValueError as e:
    print(f"Error: {e}")
    print("Please set the required environment variable:")
    print("- DEEPSEEK_API_KEY")
    exit(1)

# Initialize ELL2A
ell2a.configure({
    "enabled": True,
    "tracking_enabled": True,
    "store": "./logdir",
    "verbose": True,
    "autocommit": True,
    "model": "deepseek-reasoner-v1",
    "default_model": "deepseek-reasoner-v1",
    "client": client,
    "temperature": 0.7,
    "base_url": "https://api.deepseek.com/v1"
})

class ResearchStage(BaseModel):
    """Base model for research stages."""
    stage: str = Field(description="The current stage of research")
    domain: str = Field(description="The research domain being studied")

class LiteratureReview(ResearchStage):
    """Literature review stage output."""
    sources: List[str] = Field(description="List of academic sources reviewed")
    key_findings: List[str] = Field(description="Key findings from the literature")
    research_gaps: List[str] = Field(description="Identified gaps in current research")

class ResearchMethodology(ResearchStage):
    """Research methodology stage output."""
    methodology_type: str = Field(description="Type of research methodology")
    research_questions: List[str] = Field(description="Research questions to be addressed")
    methods: List[str] = Field(description="Specific methods to be used")
    limitations: List[str] = Field(description="Potential limitations of the methodology")

class DataAnalysis(ResearchStage):
    """Data analysis stage output."""
    analysis_method: str = Field(description="Method used for data analysis")
    key_findings: List[str] = Field(description="Key findings from the analysis")
    implications: List[str] = Field(description="Implications of the findings")
    future_work: List[str] = Field(description="Suggestions for future research")

class PaperDraft(ResearchStage):
    """Paper writing stage output."""
    title: str = Field(description="Proposed paper title")
    abstract: str = Field(description="Paper abstract")
    sections: List[str] = Field(description="Main sections of the paper")
    key_contributions: List[str] = Field(description="Key contributions of the research")

@ell2a.with_ell2a(mode="simple")
async def conduct_literature_review(domain: str, academic_level: str) -> LiteratureReview:
    """Conduct a literature review for the given domain."""
    try:
        response = client.chat.completions.create(
            model="deepseek-reasoner-v1",
            messages=[
                {
                    "role": "system",
                    "content": f"You are a {academic_level}-level research assistant conducting a literature review."
                },
                {
                    "role": "user",
                    "content": f"""Conduct a comprehensive literature review in the domain of {domain}.
Focus on recent developments, key findings, and research gaps.

Respond in this exact format:
{{
    "stage": "literature_review",
    "domain": "{domain}",
    "sources": ["source1", "source2", "source3"],
    "key_findings": ["finding1", "finding2", "finding3"],
    "research_gaps": ["gap1", "gap2", "gap3"]
}}"""
                }
            ],
            temperature=0.7
        )
        
        if not response or not response.choices or not response.choices[0].message:
            raise ValueError("No valid response received from the model")
            
        content = response.choices[0].message.content
        if not content:
            raise ValueError("Empty response content from the model")
            
        content = content.strip()
        return LiteratureReview.model_validate_json(content)
        
    except Exception as e:
        print(f"Error in literature review: {str(e)}")
        return LiteratureReview(
            stage="literature_review",
            domain=domain,
            sources=["Error retrieving sources"],
            key_findings=["Error analyzing findings"],
            research_gaps=["Error identifying gaps"]
        )

@ell2a.with_ell2a(mode="simple")
async def develop_methodology(domain: str, academic_level: str, lit_review: LiteratureReview) -> ResearchMethodology:
    """Develop research methodology based on literature review."""
    try:
        response = client.chat.completions.create(
            model="deepseek-reasoner-v1",
            messages=[
                {
                    "role": "system",
                    "content": f"You are a {academic_level}-level research assistant developing a research methodology."
                },
                {
                    "role": "user",
                    "content": f"""Based on the literature review findings, develop a research methodology for {domain}.

Literature Review Findings:
- Sources: {lit_review.sources}
- Key Findings: {lit_review.key_findings}
- Research Gaps: {lit_review.research_gaps}

Respond in this exact format:
{{
    "stage": "methodology",
    "domain": "{domain}",
    "methodology_type": "type_of_methodology",
    "research_questions": ["question1", "question2", "question3"],
    "methods": ["method1", "method2", "method3"],
    "limitations": ["limitation1", "limitation2", "limitation3"]
}}"""
                }
            ],
            temperature=0.7
        )
        
        if not response or not response.choices or not response.choices[0].message:
            raise ValueError("No valid response received from the model")
            
        content = response.choices[0].message.content
        if not content:
            raise ValueError("Empty response content from the model")
            
        content = content.strip()
        return ResearchMethodology.model_validate_json(content)
        
    except Exception as e:
        print(f"Error in methodology development: {str(e)}")
        return ResearchMethodology(
            stage="methodology",
            domain=domain,
            methodology_type="Mixed Methods Research",
            research_questions=[
                f"How can {domain} be effectively implemented in real-world settings?",
                f"What are the key factors influencing the success of {domain} applications?",
                f"How can the identified research gaps in {domain} be addressed?"
            ],
            methods=[
                "Literature analysis and systematic review",
                "Qualitative interviews with domain experts",
                "Quantitative data analysis of implementation outcomes"
            ],
            limitations=[
                "Time and resource constraints",
                "Limited access to comprehensive data",
                "Potential bias in expert selection"
            ]
        )

@ell2a.with_ell2a(mode="simple")
async def analyze_data(domain: str, academic_level: str, methodology: ResearchMethodology) -> DataAnalysis:
    """Simulate data analysis based on methodology."""
    try:
        response = client.chat.completions.create(
            model="deepseek-reasoner-v1",
            messages=[
                {
                    "role": "system",
                    "content": f"You are a {academic_level}-level research assistant conducting data analysis."
                },
                {
                    "role": "user",
                    "content": f"""Based on the research methodology, simulate a data analysis for {domain}.

Methodology:
- Type: {methodology.methodology_type}
- Research Questions: {methodology.research_questions}
- Methods: {methodology.methods}

Respond in this exact format:
{{
    "stage": "data_analysis",
    "domain": "{domain}",
    "analysis_method": "method_used",
    "key_findings": ["finding1", "finding2", "finding3"],
    "implications": ["implication1", "implication2", "implication3"],
    "future_work": ["suggestion1", "suggestion2", "suggestion3"]
}}"""
                }
            ],
            temperature=0.7
        )
        
        if not response or not response.choices or not response.choices[0].message:
            raise ValueError("No valid response received from the model")
            
        content = response.choices[0].message.content
        if not content:
            raise ValueError("Empty response content from the model")
            
        content = content.strip()
        return DataAnalysis.model_validate_json(content)
        
    except Exception as e:
        print(f"Error in data analysis: {str(e)}")
        return DataAnalysis(
            stage="data_analysis",
            domain=domain,
            analysis_method="Mixed Methods Analysis",
            key_findings=[
                f"Implementation of {domain} shows promising results in controlled environments",
                f"Success factors include proper planning, stakeholder engagement, and continuous monitoring",
                f"Challenges identified include resource constraints and resistance to change"
            ],
            implications=[
                f"Careful consideration needed when implementing {domain} solutions",
                "Stakeholder engagement is crucial for success",
                "Regular monitoring and adjustment of strategies is recommended"
            ],
            future_work=[
                f"Conduct larger-scale studies on {domain} implementation",
                "Develop standardized frameworks for evaluation",
                "Investigate long-term impacts and sustainability"
            ]
        )

@ell2a.with_ell2a(mode="simple")
async def write_paper(domain: str, academic_level: str, analysis: DataAnalysis) -> PaperDraft:
    """Write a paper draft based on the analysis."""
    try:
        response = client.chat.completions.create(
            model="deepseek-reasoner-v1",
            messages=[
                {
                    "role": "system",
                    "content": f"You are a {academic_level}-level research assistant writing an academic paper."
                },
                {
                    "role": "user",
                    "content": f"""Based on the data analysis, write a paper draft for research in {domain}.

Analysis Results:
- Method: {analysis.analysis_method}
- Key Findings: {analysis.key_findings}
- Implications: {analysis.implications}
- Future Work: {analysis.future_work}

Respond in this exact format:
{{
    "stage": "paper_writing",
    "domain": "{domain}",
    "title": "paper_title",
    "abstract": "paper_abstract",
    "sections": ["section1", "section2", "section3"],
    "key_contributions": ["contribution1", "contribution2", "contribution3"]
}}"""
                }
            ],
            temperature=0.7
        )
        
        if not response or not response.choices or not response.choices[0].message:
            raise ValueError("No valid response received from the model")
            
        content = response.choices[0].message.content
        if not content:
            raise ValueError("Empty response content from the model")
            
        content = content.strip()
        return PaperDraft.model_validate_json(content)
        
    except Exception as e:
        print(f"Error in paper writing: {str(e)}")
        return PaperDraft(
            stage="paper_writing",
            domain=domain,
            title=f"Research on {domain}",
            abstract=f"This paper presents a comprehensive study in the field of {domain}. The research explores key findings, implications, and future directions based on systematic analysis.",
            sections=["Introduction", "Literature Review", "Methodology", "Results", "Discussion", "Conclusion"],
            key_contributions=["Systematic analysis of the domain", "Identification of research gaps", "Recommendations for future work"]
        )

class AcademicAgent:
    """Academic research workflow agent."""
    
    def __init__(self, config_path: str, agent_path: str):
        """Initialize the academic agent with configuration."""
        self.config_path = config_path
        self.agent_path = agent_path
        self.load_configurations()
        
    def load_configurations(self):
        """Load configuration files."""
        try:
            with open(self.config_path, 'r') as f:
                self.config = json.load(f)
            with open(self.agent_path, 'r') as f:
                self.agent_config = json.load(f)
        except Exception as e:
            print(f"Error loading configurations: {e}")
            raise
            
    def execute_workflow(self, input_data: Dict[str, Any]) -> Dict[str, Any]:
        """Execute the research workflow."""
        try:
            # Extract input parameters
            needs = input_data.get('student_needs', {})
            domain = needs.get('research_topic', '')
            academic_level = needs.get('academic_level', '')
            
            # Run the workflow
            results = asyncio.run(self._run_workflow(domain, academic_level))
            return results
            
        except Exception as e:
            print(f"Error executing workflow: {e}")
            raise
            
    async def _run_workflow(self, domain: str, academic_level: str) -> Dict[str, Any]:
        """Run the research workflow asynchronously."""
        try:
            # Step 1: Literature Review
            print("Conducting literature review...")
            lit_review = await conduct_literature_review(domain, academic_level)
            
            # Step 2: Research Methodology
            print("Developing methodology...")
            methodology = await develop_methodology(domain, academic_level, lit_review)
            
            # Step 3: Data Analysis
            print("Analyzing data...")
            analysis = await analyze_data(domain, academic_level, methodology)
            
            # Step 4: Paper Writing
            print("Writing paper...")
            paper = await write_paper(domain, academic_level, analysis)
            
            return {
                'literature_review': lit_review.model_dump(),
                'methodology': methodology.model_dump(),
                'analysis': analysis.model_dump(),
                'paper': paper.model_dump()
            }
            
        except Exception as e:
            print(f"Error in workflow execution: {e}")
            raise
            
    def generate_output_document(self, results: Dict[str, Any], format_type: str, output_path: str) -> str:
        """Generate output document in specified format."""
        try:
            # Create output directory if it doesn't exist
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Basic markdown content generation
            content = f"""# Research Plan: {results['paper']['title']}

## Abstract
{results['paper']['abstract']}

## Literature Review
### Key Findings
{chr(10).join('- ' + finding for finding in results['literature_review']['key_findings'])}

### Research Gaps
{chr(10).join('- ' + gap for gap in results['literature_review']['research_gaps'])}

## Methodology
### Research Questions
{chr(10).join('- ' + question for question in results['methodology']['research_questions'])}

### Methods
{chr(10).join('- ' + method for method in results['methodology']['methods'])}

## Analysis
### Key Findings
{chr(10).join('- ' + finding for finding in results['analysis']['key_findings'])}

### Implications
{chr(10).join('- ' + impl for impl in results['analysis']['implications'])}

## Future Work
{chr(10).join('- ' + work for work in results['analysis']['future_work'])}
"""
            
            # Write content based on format
            if format_type == 'markdown':
                with open(output_path, 'w') as f:
                    f.write(content)
                
            elif format_type == 'docx':
                # Create a new Word document
                doc = Document()
                
                # Add title
                doc.add_heading(f"Research Plan: {results['paper']['title']}", 0)
                
                # Add abstract
                doc.add_heading('Abstract', level=1)
                doc.add_paragraph(results['paper']['abstract'])
                
                # Add Literature Review
                doc.add_heading('Literature Review', level=1)
                doc.add_heading('Key Findings', level=2)
                for finding in results['literature_review']['key_findings']:
                    doc.add_paragraph(finding, style='List Bullet')
                    
                doc.add_heading('Research Gaps', level=2)
                for gap in results['literature_review']['research_gaps']:
                    doc.add_paragraph(gap, style='List Bullet')
                    
                # Add Methodology
                doc.add_heading('Methodology', level=1)
                doc.add_heading('Research Questions', level=2)
                for question in results['methodology']['research_questions']:
                    doc.add_paragraph(question, style='List Bullet')
                    
                doc.add_heading('Methods', level=2)
                for method in results['methodology']['methods']:
                    doc.add_paragraph(method, style='List Bullet')
                    
                # Add Analysis
                doc.add_heading('Analysis', level=1)
                doc.add_heading('Key Findings', level=2)
                for finding in results['analysis']['key_findings']:
                    doc.add_paragraph(finding, style='List Bullet')
                    
                doc.add_heading('Implications', level=2)
                for impl in results['analysis']['implications']:
                    doc.add_paragraph(impl, style='List Bullet')
                    
                # Add Future Work
                doc.add_heading('Future Work', level=1)
                for work in results['analysis']['future_work']:
                    doc.add_paragraph(work, style='List Bullet')
                    
                # Save the document
                doc.save(output_path)
                
            elif format_type == 'pdf':
                # Create PDF using reportlab
                doc = SimpleDocTemplate(output_path, pagesize=letter)
                styles = getSampleStyleSheet()
                story = []
                
                # Add title
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Title'],
                    fontSize=24,
                    spaceAfter=30
                )
                story.append(Paragraph(f"Research Plan: {results['paper']['title']}", title_style))
                story.append(Spacer(1, 12))
                
                # Add abstract
                story.append(Paragraph('Abstract', styles['Heading1']))
                story.append(Paragraph(results['paper']['abstract'], styles['Normal']))
                story.append(Spacer(1, 12))
                
                # Add Literature Review
                story.append(Paragraph('Literature Review', styles['Heading1']))
                story.append(Paragraph('Key Findings', styles['Heading2']))
                for finding in results['literature_review']['key_findings']:
                    story.append(Paragraph(f"• {finding}", styles['Normal']))
                
                story.append(Paragraph('Research Gaps', styles['Heading2']))
                for gap in results['literature_review']['research_gaps']:
                    story.append(Paragraph(f"• {gap}", styles['Normal']))
                
                # Add Methodology
                story.append(Paragraph('Methodology', styles['Heading1']))
                story.append(Paragraph('Research Questions', styles['Heading2']))
                for question in results['methodology']['research_questions']:
                    story.append(Paragraph(f"• {question}", styles['Normal']))
                
                story.append(Paragraph('Methods', styles['Heading2']))
                for method in results['methodology']['methods']:
                    story.append(Paragraph(f"• {method}", styles['Normal']))
                
                # Add Analysis
                story.append(Paragraph('Analysis', styles['Heading1']))
                story.append(Paragraph('Key Findings', styles['Heading2']))
                for finding in results['analysis']['key_findings']:
                    story.append(Paragraph(f"• {finding}", styles['Normal']))
                
                story.append(Paragraph('Implications', styles['Heading2']))
                for impl in results['analysis']['implications']:
                    story.append(Paragraph(f"• {impl}", styles['Normal']))
                
                # Add Future Work
                story.append(Paragraph('Future Work', styles['Heading1']))
                for work in results['analysis']['future_work']:
                    story.append(Paragraph(f"• {work}", styles['Normal']))
                
                # Build PDF
                doc.build(story)
                
            return output_path
            
        except Exception as e:
            print(f"Error generating output document: {e}")
            raise

async def main():
    """Main function to demonstrate the academic research workflow."""
    # Research parameters
    domain = "Quantum Machine Learning"
    academic_level = "PhD"
    
    print(f"\nStarting {academic_level}-level research on {domain}")
    print("=" * 50)
    
    # Step 1: Literature Review
    print("\n1. Conducting Literature Review...")
    lit_review = await conduct_literature_review(domain, academic_level)
    print("\nKey Findings:")
    for i, finding in enumerate(lit_review.key_findings, 1):
        print(f"{i}. {finding}")
    print("\nResearch Gaps:")
    for i, gap in enumerate(lit_review.research_gaps, 1):
        print(f"{i}. {gap}")
    
    # Step 2: Research Methodology
    print("\n2. Developing Research Methodology...")
    methodology = await develop_methodology(domain, academic_level, lit_review)
    print("\nResearch Questions:")
    for i, question in enumerate(methodology.research_questions, 1):
        print(f"{i}. {question}")
    print("\nMethods:")
    for i, method in enumerate(methodology.methods, 1):
        print(f"{i}. {method}")
    
    # Step 3: Data Analysis
    print("\n3. Conducting Data Analysis...")
    analysis = await analyze_data(domain, academic_level, methodology)
    print("\nKey Findings:")
    for i, finding in enumerate(analysis.key_findings, 1):
        print(f"{i}. {finding}")
    print("\nImplications:")
    for i, implication in enumerate(analysis.implications, 1):
        print(f"{i}. {implication}")
    
    # Step 4: Paper Writing
    print("\n4. Writing Paper Draft...")
    paper = await write_paper(domain, academic_level, analysis)
    print(f"\nTitle: {paper.title}")
    print(f"\nAbstract: {paper.abstract}")
    print("\nKey Contributions:")
    for i, contribution in enumerate(paper.key_contributions, 1):
        print(f"{i}. {contribution}")

if __name__ == "__main__":
    asyncio.run(main())